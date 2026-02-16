#!/usr/bin/env python3
"""
Word Document Generator for ADB Framework Documentation
Converts Markdown documents to professional Word format
"""

import os
import sys
from pathlib import Path
import re
from datetime import datetime

def install_dependencies():
    """Install required dependencies"""
    try:
        import python_docx
        import markdown
    except ImportError:
        import subprocess
        print("Installing required packages...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "markdown"])

def generate_word_docs():
    """Generate Word documents from Markdown files"""
    
    # Install dependencies if needed
    install_dependencies()
    
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    import markdown
    
    # Files to convert
    files_to_convert = [
        {
            'input': 'IMPLEMENTATION_REVIEW_v1.0.0.md',
            'output': 'ADB_Framework_Implementation_Review_v1.0.0.docx',
            'title': 'ADB Framework Telco Automation - Implementation Review v1.0.0'
        },
        {
            'input': 'RELEASE_NOTES_v1.0.0_REVISED.md',
            'output': 'ADB_Framework_Release_Notes_v1.0.0_Revised.docx',
            'title': 'ADB Framework Telco Automation - Release Notes v1.0.0 (Revised)'
        }
    ]
    
    for file_info in files_to_convert:
        input_file = Path(file_info['input'])
        output_file = Path(file_info['output'])
        
        if not input_file.exists():
            print(f"Warning: {input_file} not found, skipping...")
            continue
        
        print(f"Converting {input_file} to {output_file}...")
        
        # Read Markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = file_info['title']
        doc.core_properties.author = "F2G Telco Academy"
        doc.core_properties.subject = "ADB Framework Documentation"
        doc.core_properties.created = datetime.now()
        
        # Define styles
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = None  # Default color
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(18)
        
        # Heading 1 style
        heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_font = heading1_style.font
        heading1_font.name = 'Calibri'
        heading1_font.size = Pt(18)
        heading1_font.bold = True
        heading1_style.paragraph_format.space_before = Pt(12)
        heading1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_font = heading2_style.font
        heading2_font.name = 'Calibri'
        heading2_font.size = Pt(14)
        heading2_font.bold = True
        heading2_style.paragraph_format.space_before = Pt(10)
        heading2_style.paragraph_format.space_after = Pt(4)
        
        # Heading 3 style
        heading3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        heading3_font = heading3_style.font
        heading3_font.name = 'Calibri'
        heading3_font.size = Pt(12)
        heading3_font.bold = True
        heading3_style.paragraph_format.space_before = Pt(8)
        heading3_style.paragraph_format.space_after = Pt(4)
        
        # Body style
        body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Code style
        code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Process markdown content
        lines = md_content.split('\n')
        current_paragraph = []
        in_code_block = False
        code_block = []
        
        # Add title page
        title_para = doc.add_paragraph()
        title_para.style = title_style
        title_para.add_run(file_info['title'])
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)
        date_run.font.italic = True
        
        # Add page break
        doc.add_page_break()
        
        for line in lines:
            line = line.strip()
            
            if not line and not in_code_block:
                if current_paragraph:
                    para = doc.add_paragraph()
                    para.style = body_style
                    para_text = ' '.join(current_paragraph)
                    # Remove emojis and clean text
                    para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                    para.add_run(para_text)
                    current_paragraph = []
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_block:
                        code_para = doc.add_paragraph()
                        code_para.style = code_style
                        code_text = '\n'.join(code_block)
                        code_para.add_run(code_text)
                        code_block = []
                    in_code_block = False
                else:
                    # Start code block
                    if current_paragraph:
                        para = doc.add_paragraph()
                        para.style = body_style
                        para_text = ' '.join(current_paragraph)
                        para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                        para.add_run(para_text)
                        current_paragraph = []
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block.append(line)
                continue
            
            # Handle headers
            if line.startswith('# '):
                if current_paragraph:
                    para = doc.add_paragraph()
                    para.style = body_style
                    para_text = ' '.join(current_paragraph)
                    para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                    para.add_run(para_text)
                    current_paragraph = []
                
                header_text = line[2:].strip()
                header_text = re.sub(r'[^\w\s\-\.\(\)]', '', header_text)
                header_para = doc.add_paragraph()
                header_para.style = title_style
                header_para.add_run(header_text)
                
            elif line.startswith('## '):
                if current_paragraph:
                    para = doc.add_paragraph()
                    para.style = body_style
                    para_text = ' '.join(current_paragraph)
                    para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                    para.add_run(para_text)
                    current_paragraph = []
                
                header_text = line[3:].strip()
                header_text = re.sub(r'[^\w\s\-\.\(\)]', '', header_text)
                header_para = doc.add_paragraph()
                header_para.style = heading1_style
                header_para.add_run(header_text)
                
            elif line.startswith('### '):
                if current_paragraph:
                    para = doc.add_paragraph()
                    para.style = body_style
                    para_text = ' '.join(current_paragraph)
                    para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                    para.add_run(para_text)
                    current_paragraph = []
                
                header_text = line[4:].strip()
                header_text = re.sub(r'[^\w\s\-\.\(\)]', '', header_text)
                header_para = doc.add_paragraph()
                header_para.style = heading2_style
                header_para.add_run(header_text)
                
            elif line.startswith('#### '):
                if current_paragraph:
                    para = doc.add_paragraph()
                    para.style = body_style
                    para_text = ' '.join(current_paragraph)
                    para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                    para.add_run(para_text)
                    current_paragraph = []
                
                header_text = line[5:].strip()
                header_text = re.sub(r'[^\w\s\-\.\(\)]', '', header_text)
                header_para = doc.add_paragraph()
                header_para.style = heading3_style
                header_para.add_run(header_text)
                
            elif line.startswith('- ') or line.startswith('* '):
                if current_paragraph:
                    para = doc.add_paragraph()
                    para.style = body_style
                    para_text = ' '.join(current_paragraph)
                    para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                    para.add_run(para_text)
                    current_paragraph = []
                
                bullet_text = line[2:].strip()
                bullet_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', bullet_text)
                bullet_para = doc.add_paragraph()
                bullet_para.style = body_style
                bullet_para.add_run(f'• {bullet_text}')
                
            elif line.startswith('---'):
                if current_paragraph:
                    para = doc.add_paragraph()
                    para.style = body_style
                    para_text = ' '.join(current_paragraph)
                    para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
                    para.add_run(para_text)
                    current_paragraph = []
                # Add horizontal line (paragraph break)
                doc.add_paragraph()
                
            else:
                # Regular text
                if line:
                    clean_line = re.sub(r'[^\w\s\-\.\(\):/,]', '', line)
                    if clean_line.strip():
                        current_paragraph.append(clean_line)
        
        # Add remaining content
        if current_paragraph:
            para = doc.add_paragraph()
            para.style = body_style
            para_text = ' '.join(current_paragraph)
            para_text = re.sub(r'[^\w\s\-\.\(\):/,]', '', para_text)
            para.add_run(para_text)
        
        # Save document
        try:
            doc.save(output_file)
            print(f"Successfully created: {output_file}")
            print(f"File size: {output_file.stat().st_size / 1024:.1f} KB")
        except Exception as e:
            print(f"Error saving {output_file}: {e}")
            return False
    
    return True

if __name__ == "__main__":
    print("ADB Framework Word Document Generator")
    print("=" * 50)
    
    success = generate_word_docs()
    
    if success:
        print("\nWord document generation completed successfully!")
        print("You can find the Word files in the current directory.")
    else:
        print("\nWord document generation failed!")
        sys.exit(1)